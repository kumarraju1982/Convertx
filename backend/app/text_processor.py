"""
Text post-processor for fixing common PDF conversion issues.

This module provides utilities to fix text extracted from PDFs,
such as missing spaces between words.
"""

import re
from docx import Document
from typing import List


class TextProcessor:
    """
    Post-processes converted Word documents to fix common issues.
    """
    
    def __init__(self):
        """Initialize the text processor."""
        pass
    
    def add_spaces_to_text(self, text: str) -> str:
        """
        Add spaces between words that are incorrectly joined.
        
        Uses heuristics to detect word boundaries:
        - Lowercase followed by uppercase (camelCase)
        - Letter followed by number or vice versa
        - Multiple capital letters followed by lowercase
        
        Args:
            text: Input text with missing spaces
            
        Returns:
            Text with spaces added
        """
        if not text:
            return text
        
        # Add space between lowercase and uppercase (camelCase)
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Add space between letter and number
        text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)
        text = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', text)
        
        # Add space between multiple capitals and lowercase
        # e.g., "RFPis" -> "RFP is"
        text = re.sub(r'([A-Z]{2,})([a-z])', r'\1 \2', text)
        
        # Add space after punctuation if missing
        text = re.sub(r'([.!?,;:])([A-Za-z])', r'\1 \2', text)
        
        # Fix common word patterns (add more as needed)
        common_fixes = {
            'tothis': 'to this',
            'tothe': 'to the',
            'ofthe': 'of the',
            'inthe': 'in the',
            'onthe': 'on the',
            'forthe': 'for the',
            'andthe': 'and the',
            'withthe': 'with the',
            'fromthe': 'from the',
            'aboutthe': 'about the',
            'thatis': 'that is',
            'whichis': 'which is',
            'thereis': 'there is',
            'itis': 'it is',
        }
        
        for wrong, correct in common_fixes.items():
            text = re.sub(r'\b' + wrong + r'\b', correct, text, flags=re.IGNORECASE)
        
        return text
    
    def fix_word_document(self, doc_path: str) -> None:
        """
        Fix spacing issues in a Word document.
        
        Opens the document, processes all paragraphs and table cells,
        and saves the fixed version.
        
        Args:
            doc_path: Path to the Word document
        """
        try:
            # Open document
            doc = Document(doc_path)
            
            # Fix paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    # Store original formatting
                    runs = paragraph.runs
                    if runs:
                        # Get text from all runs
                        original_text = paragraph.text
                        fixed_text = self.add_spaces_to_text(original_text)
                        
                        if fixed_text != original_text:
                            # Clear paragraph and add fixed text
                            paragraph.clear()
                            paragraph.add_run(fixed_text)
            
            # Fix tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text:
                                original_text = paragraph.text
                                fixed_text = self.add_spaces_to_text(original_text)
                                
                                if fixed_text != original_text:
                                    paragraph.clear()
                                    paragraph.add_run(fixed_text)
            
            # Save document
            doc.save(doc_path)
            
        except Exception as e:
            print(f"Error fixing Word document: {e}")
            # Don't raise - better to have document with spacing issues than no document
