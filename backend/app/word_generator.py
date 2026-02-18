"""
Word Generator for creating Microsoft Word documents from structured content.

This module converts DocumentStructure objects into formatted Word documents
using the python-docx library, preserving formatting for paragraphs, headings,
lists, and tables.
"""

from typing import List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.models import DocumentStructure, StructureElement


class WordGenerator:
    """
    Generates Microsoft Word documents from structured content.
    
    This class converts DocumentStructure objects (containing paragraphs,
    headings, lists, and tables) into formatted .docx files using python-docx.
    """
    
    def __init__(self):
        """Initialize the Word Generator."""
        pass
    
    def create_document(self, structures: List[DocumentStructure]) -> Document:
        """
        Create a Word document from document structures.
        
        Processes a list of DocumentStructure objects (typically one per page)
        and creates a formatted Word document with appropriate styling for
        paragraphs, headings, lists, and tables.
        
        Args:
            structures: List of DocumentStructure objects to convert
            
        Returns:
            Document object from python-docx
            
        Requirements:
            - 4.1: Create a valid .docx file
            - 4.2: Apply detected formatting (headings, paragraphs, lists)
            - 4.3: Create Word table structures
        """
        doc = Document()
        
        # Process each page structure
        for page_idx, structure in enumerate(structures):
            # Add page break between pages (except before first page)
            if page_idx > 0:
                doc.add_page_break()
            
            # Process each element in the structure
            for element in structure.elements:
                self._add_element_to_document(doc, element)
        
        return doc
    
    def _add_element_to_document(self, doc: Document, element: StructureElement) -> None:
        """
        Add a structure element to the document with appropriate formatting.
        
        Args:
            doc: Document object to add element to
            element: StructureElement to add
        """
        if element.type == "heading":
            self._add_heading(doc, element)
        elif element.type == "paragraph":
            self._add_paragraph(doc, element)
        elif element.type == "list":
            self._add_list(doc, element)
        elif element.type == "table":
            self._add_table(doc, element)
    
    def _add_heading(self, doc: Document, element: StructureElement) -> None:
        """
        Add a heading to the document.
        
        Args:
            doc: Document object
            element: StructureElement with type="heading"
            
        Requirements:
            - 4.2: Apply detected formatting (headings)
        """
        # Ensure level is within valid range (1-3 as per design)
        level = max(1, min(3, element.level))
        
        # Add heading with appropriate level
        doc.add_heading(element.content, level=level)
    
    def _add_paragraph(self, doc: Document, element: StructureElement) -> None:
        """
        Add a paragraph to the document.
        
        Args:
            doc: Document object
            element: StructureElement with type="paragraph"
            
        Requirements:
            - 4.2: Apply detected formatting (paragraphs)
        """
        # Split content by newlines to handle multi-line paragraphs
        lines = element.content.split('\n')
        
        for line in lines:
            if line.strip():  # Only add non-empty lines
                doc.add_paragraph(line.strip())
    
    def _add_list(self, doc: Document, element: StructureElement) -> None:
        """
        Add a list to the document.
        
        Args:
            doc: Document object
            element: StructureElement with type="list"
            
        Requirements:
            - 4.2: Apply detected formatting (lists)
        """
        # Get list type from style
        list_type = element.style.get("list_type", "bullet")
        
        # Split content into individual list items
        items = element.content.split('\n')
        
        for item in items:
            if item.strip():
                # Add list item as a paragraph with bullet/number style
                # python-docx uses 'List Bullet' and 'List Number' styles
                if list_type == "numbered":
                    doc.add_paragraph(item.strip(), style='List Number')
                else:
                    doc.add_paragraph(item.strip(), style='List Bullet')
    
    def _add_table(self, doc: Document, element: StructureElement) -> None:
        """
        Add a table to the document.
        
        Args:
            doc: Document object
            element: StructureElement with type="table"
            
        Requirements:
            - 4.3: Create Word table structures
        """
        # Parse table content (rows separated by newlines)
        rows_text = element.content.split('\n')
        rows_text = [row for row in rows_text if row.strip()]
        
        if not rows_text:
            return
        
        # Get table dimensions from style
        num_rows = element.style.get("rows", len(rows_text))
        num_cols = element.style.get("columns", 2)  # Default to 2 columns
        
        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Light Grid Accent 1'
        
        # Fill table with content
        # Split each row by whitespace to distribute across columns
        for row_idx, row_text in enumerate(rows_text):
            if row_idx >= num_rows:
                break
            
            # Split row text into cells
            # Use multiple spaces as delimiter to separate columns
            cells_text = row_text.split()
            
            # Distribute text across columns
            row_cells = table.rows[row_idx].cells
            
            # Strategy: distribute words evenly across columns
            words_per_col = len(cells_text) // num_cols
            remainder = len(cells_text) % num_cols
            
            word_idx = 0
            for col_idx in range(num_cols):
                # Calculate how many words for this column
                col_word_count = words_per_col + (1 if col_idx < remainder else 0)
                
                # Get words for this column
                col_words = cells_text[word_idx:word_idx + col_word_count]
                word_idx += col_word_count
                
                # Set cell text
                if col_words:
                    row_cells[col_idx].text = ' '.join(col_words)
    
    def save(self, document: Document, output_path: str, overwrite: bool = True) -> bool:
        """
        Save the document to the specified path with validation.

        Args:
            document: Document object to save
            output_path: Path where the document should be saved
            overwrite: If False and file exists, generate unique filename

        Returns:
            True if save was successful, False otherwise

        Raises:
            FileIOError: If output directory doesn't exist or path is invalid

        Requirements:
            - 4.4: Save document to specified output path
            - 4.5: Validate output path before saving
            - 8.2: Validate directory exists
            - 8.3: Handle file conflicts
        """
        import os
        from app.exceptions import FileIOError

        try:
            # Validate output path
            output_dir = os.path.dirname(output_path)

            # If no directory specified, use current directory
            if not output_dir:
                output_dir = '.'

            # Check if directory exists
            if not os.path.exists(output_dir):
                raise FileIOError(f"Output directory does not exist: {output_dir}")

            # Check if directory is writable
            if not os.access(output_dir, os.W_OK):
                raise FileIOError(f"Output directory is not writable: {output_dir}")

            # Handle file conflicts
            final_path = output_path
            if os.path.exists(output_path) and not overwrite:
                # Generate unique filename
                base, ext = os.path.splitext(output_path)
                counter = 1
                while os.path.exists(final_path):
                    final_path = f"{base}_{counter}{ext}"
                    counter += 1

            # Save the document
            document.save(final_path)
            return True

        except FileIOError:
            # Re-raise FileIOError as-is
            raise
        except Exception as e:
            # Wrap other exceptions in FileIOError
            raise FileIOError(f"Error saving document: {e}")


