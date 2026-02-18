"""
Unit tests for WordGenerator class.

Tests the conversion of DocumentStructure objects into Word documents,
including formatting for paragraphs, headings, lists, and tables.
"""

import pytest
import os
import tempfile
from docx.document import Document as DocxDocument
from app.word_generator import WordGenerator
from app.models import DocumentStructure, StructureElement


class TestWordGenerator:
    """Test suite for WordGenerator class."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.generator = WordGenerator()
    
    def test_create_empty_document(self):
        """Test creating a document from empty structures."""
        structures = [DocumentStructure(elements=[])]
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        assert isinstance(doc, DocxDocument)
    
    def test_create_document_with_paragraph(self):
        """Test creating a document with a simple paragraph."""
        element = StructureElement(
            type="paragraph",
            content="This is a test paragraph.",
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        assert len(doc.paragraphs) > 0
        assert "This is a test paragraph." in doc.paragraphs[0].text
    
    def test_create_document_with_heading_level_1(self):
        """Test creating a document with a level 1 heading."""
        element = StructureElement(
            type="heading",
            content="Main Heading",
            level=1,
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        assert len(doc.paragraphs) > 0
        assert "Main Heading" in doc.paragraphs[0].text
        # Check that it's styled as a heading
        assert doc.paragraphs[0].style.name.startswith('Heading')
    
    def test_create_document_with_heading_level_2(self):
        """Test creating a document with a level 2 heading."""
        element = StructureElement(
            type="heading",
            content="Subheading",
            level=2,
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        assert "Subheading" in doc.paragraphs[0].text
        assert doc.paragraphs[0].style.name == 'Heading 2'
    
    def test_create_document_with_heading_level_3(self):
        """Test creating a document with a level 3 heading."""
        element = StructureElement(
            type="heading",
            content="Sub-subheading",
            level=3,
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        assert "Sub-subheading" in doc.paragraphs[0].text
        assert doc.paragraphs[0].style.name == 'Heading 3'
    
    def test_create_document_with_bullet_list(self):
        """Test creating a document with a bullet list."""
        element = StructureElement(
            type="list",
            content="• First item\n• Second item\n• Third item",
            style={"list_type": "bullet"}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        # Check that list items were added
        list_paragraphs = [p for p in doc.paragraphs if 'List' in p.style.name]
        assert len(list_paragraphs) == 3
        assert "First item" in list_paragraphs[0].text
        assert "Second item" in list_paragraphs[1].text
        assert "Third item" in list_paragraphs[2].text
    
    def test_create_document_with_numbered_list(self):
        """Test creating a document with a numbered list."""
        element = StructureElement(
            type="list",
            content="1. First item\n2. Second item\n3. Third item",
            style={"list_type": "numbered"}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        # Check that numbered list items were added
        list_paragraphs = [p for p in doc.paragraphs if 'List Number' in p.style.name]
        assert len(list_paragraphs) == 3
    
    def test_create_document_with_table(self):
        """Test creating a document with a table."""
        element = StructureElement(
            type="table",
            content="Header1 Header2\nRow1Col1 Row1Col2\nRow2Col1 Row2Col2",
            style={"rows": 3, "columns": 2}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3
        assert len(table.columns) == 2
    
    def test_create_document_with_mixed_elements(self):
        """Test creating a document with multiple element types."""
        elements = [
            StructureElement(type="heading", content="Document Title", level=1, style={}),
            StructureElement(type="paragraph", content="Introduction paragraph.", style={}),
            StructureElement(type="heading", content="Section 1", level=2, style={}),
            StructureElement(type="list", content="• Item 1\n• Item 2", style={"list_type": "bullet"}),
            StructureElement(type="table", content="A B\nC D", style={"rows": 2, "columns": 2})
        ]
        structures = [DocumentStructure(elements=elements)]
        
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        assert len(doc.paragraphs) >= 5  # Headings, paragraph, and list items
        assert len(doc.tables) == 1
    
    def test_create_document_with_multiple_pages(self):
        """Test creating a document from multiple page structures."""
        page1 = DocumentStructure(elements=[
            StructureElement(type="heading", content="Page 1", level=1, style={}),
            StructureElement(type="paragraph", content="Content of page 1.", style={})
        ])
        page2 = DocumentStructure(elements=[
            StructureElement(type="heading", content="Page 2", level=1, style={}),
            StructureElement(type="paragraph", content="Content of page 2.", style={})
        ])
        
        doc = self.generator.create_document([page1, page2])
        
        assert doc is not None
        # Check that both pages' content is present
        all_text = ' '.join(p.text for p in doc.paragraphs)
        assert "Page 1" in all_text
        assert "Page 2" in all_text
        assert "Content of page 1" in all_text
        assert "Content of page 2" in all_text
    
    def test_create_document_with_multiline_paragraph(self):
        """Test creating a document with a multi-line paragraph."""
        element = StructureElement(
            type="paragraph",
            content="Line 1\nLine 2\nLine 3",
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        assert doc is not None
        # Each line should become a separate paragraph
        assert len(doc.paragraphs) >= 3
    
    def test_save_document_success(self):
        """Test saving a document to a file."""
        element = StructureElement(
            type="paragraph",
            content="Test content",
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        doc = self.generator.create_document(structures)
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Remove the file so we can test creation
            os.remove(tmp_path)
            
            # Save the document
            result = self.generator.save(doc, tmp_path)
            
            assert result is True
            assert os.path.exists(tmp_path)
            assert os.path.getsize(tmp_path) > 0
            
            # Verify the file can be opened as a Word document
            from docx import Document
            saved_doc = Document(tmp_path)
            assert saved_doc is not None
        finally:
            # Clean up
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    def test_save_document_invalid_path(self):
        """Test saving a document to an invalid path."""
        from app.exceptions import FileIOError
        
        element = StructureElement(
            type="paragraph",
            content="Test content",
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        doc = self.generator.create_document(structures)
        
        # Try to save to an invalid path (non-existent directory)
        invalid_path = "/nonexistent/directory/file.docx"
        
        with pytest.raises(FileIOError) as exc_info:
            self.generator.save(doc, invalid_path)
        
        assert "does not exist" in str(exc_info.value)
    
    def test_heading_level_clamping(self):
        """Test that heading levels are clamped to valid range (1-3)."""
        # Test level 0 (should become 1)
        element1 = StructureElement(type="heading", content="Level 0", level=0, style={})
        structures1 = [DocumentStructure(elements=[element1])]
        doc1 = self.generator.create_document(structures1)
        assert doc1.paragraphs[0].style.name == 'Heading 1'
        
        # Test level 4 (should become 3)
        element2 = StructureElement(type="heading", content="Level 4", level=4, style={})
        structures2 = [DocumentStructure(elements=[element2])]
        doc2 = self.generator.create_document(structures2)
        assert doc2.paragraphs[0].style.name == 'Heading 3'
    
    def test_empty_list_items_ignored(self):
        """Test that empty list items are ignored."""
        element = StructureElement(
            type="list",
            content="• Item 1\n\n• Item 2\n",
            style={"list_type": "bullet"}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        # Only non-empty items should be added
        list_paragraphs = [p for p in doc.paragraphs if 'List' in p.style.name]
        assert len(list_paragraphs) == 2
    
    def test_empty_paragraph_lines_ignored(self):
        """Test that empty lines in paragraphs are ignored."""
        element = StructureElement(
            type="paragraph",
            content="Line 1\n\nLine 2\n",
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        # Only non-empty lines should create paragraphs
        non_empty_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        assert len(non_empty_paragraphs) == 2
    
    def test_table_with_empty_content(self):
        """Test creating a table with empty content."""
        element = StructureElement(
            type="table",
            content="",
            style={"rows": 2, "columns": 2}
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        # Table should not be created if content is empty
        assert len(doc.tables) == 0
    
    def test_table_with_default_columns(self):
        """Test creating a table without explicit column count."""
        element = StructureElement(
            type="table",
            content="A B C\nD E F",
            style={"rows": 2}  # No columns specified
        )
        structures = [DocumentStructure(elements=[element])]
        
        doc = self.generator.create_document(structures)
        
        assert len(doc.tables) == 1
        # Should default to 2 columns
        assert len(doc.tables[0].columns) == 2
    
    def test_save_document_overwrite_existing(self):
        """Test saving a document with overwrite=True (default)."""
        element = StructureElement(
            type="paragraph",
            content="Original content",
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        doc1 = self.generator.create_document(structures)
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Save first document
            self.generator.save(doc1, tmp_path)
            first_size = os.path.getsize(tmp_path)
            
            # Create a different document
            element2 = StructureElement(
                type="paragraph",
                content="New content with much more text to make it larger",
                style={}
            )
            structures2 = [DocumentStructure(elements=[element2])]
            doc2 = self.generator.create_document(structures2)
            
            # Save second document with overwrite=True (should replace)
            result = self.generator.save(doc2, tmp_path, overwrite=True)
            
            assert result is True
            assert os.path.exists(tmp_path)
            # File should be overwritten (size might be different)
            second_size = os.path.getsize(tmp_path)
            # Just verify it's still a valid document
            from docx import Document
            saved_doc = Document(tmp_path)
            assert saved_doc is not None
        finally:
            # Clean up
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    def test_save_document_unique_filename(self):
        """Test saving a document with overwrite=False generates unique filename."""
        element = StructureElement(
            type="paragraph",
            content="Test content",
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        doc = self.generator.create_document(structures)
        
        # Create a temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Save first document
            self.generator.save(doc, tmp_path)
            
            # Save second document with overwrite=False
            # This should create a new file with _1 suffix
            result = self.generator.save(doc, tmp_path, overwrite=False)
            
            assert result is True
            
            # Original file should still exist
            assert os.path.exists(tmp_path)
            
            # New file with _1 suffix should exist
            base, ext = os.path.splitext(tmp_path)
            expected_new_path = f"{base}_1{ext}"
            assert os.path.exists(expected_new_path)
            
            # Clean up the new file
            if os.path.exists(expected_new_path):
                os.remove(expected_new_path)
        finally:
            # Clean up original file
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    def test_save_document_current_directory(self):
        """Test saving a document with no directory specified."""
        element = StructureElement(
            type="paragraph",
            content="Test content",
            style={}
        )
        structures = [DocumentStructure(elements=[element])]
        doc = self.generator.create_document(structures)
        
        # Use just a filename (no directory)
        filename = "test_output.docx"
        
        try:
            # Save the document
            result = self.generator.save(doc, filename)
            
            assert result is True
            assert os.path.exists(filename)
            
            # Verify the file can be opened
            from docx import Document
            saved_doc = Document(filename)
            assert saved_doc is not None
        finally:
            # Clean up
            if os.path.exists(filename):
                os.remove(filename)
